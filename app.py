import streamlit as st
import calendar
import io
import json
import hashlib
import re

from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


# ============================================================
# 頁面設定
# ============================================================

st.set_page_config(
    page_title="托嬰中心 教案與適性月計畫系統",
    layout="wide",
    initial_sidebar_state="collapsed",
)


# ============================================================
# Word 排版工具
# ============================================================

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=70, bottom=70, left=70, right=70):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table, color="1F497D", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)


def set_run_font(run, size=9, bold=False, color=None):
    run.font.name = "微軟正黑體"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


# ============================================================
# 適齡指標參考池
# 重要：AI只能從這裡選指標，不得自行發明代碼。
# 如中心未來補齊官方指標，可直接擴充此區，不必改AI邏輯。
# ============================================================

INDICATOR_POOLS = {
    "0-12個月": [
        "I-1-1：能在俯臥時抬頭並維持短暫姿勢",
        "II-1-1：能以手臂支撐上半身",
        "II-1-2：能翻身",
        "II-1-3：能在協助下坐起",
        "III-1-1：能自己坐穩",
        "III-1-2：會撐起身體向前爬行",
        "III-1-3：能在協助下站立與行走",
        "III-1-5：能將物品由一手換到另一手",
        "III-1-6：能用拇指配合其他手指鉗握物品",
        "III-1-7：能以雙手拍手",
        "III-2-1：會揮手表示再見",
        "III-2-2：能和成人玩簡單重複遊戲",
        "III-2-3：能簡單表達自我需求及情緒",
        "III-2-5：能與照顧者建立情感依附",
        "III-3-1：能模仿大人的簡單話語",
        "III-3-3：會以肢體動作進行溝通",
        "III-3-4：會與人輪流對話",
        "III-3-5：能理解簡單語彙的意思",
        "III-4-2：會尋找完全被藏著的物品",
        "III-4-3：能預期事件的發生",
        "III-4-4：呼叫名字時會有反應",
        "III-4-5：能設法接近想要的事物",
        "III-4-6：能操作簡單玩具",
        "III-5-1：能自己拿住奶瓶進食",
        "III-5-2：能吞嚥糊狀副食品",
        "III-5-3：能自己拿食物吃",
        "III-5-4：能拉下頭上的帽子",
        "III-5-5：會表示要吃東西",
    ],
    "7-12個月": [
        "III-1-1：能自己坐穩",
        "III-1-2：會撐起身體向前爬行",
        "III-1-3：能在協助下站立與行走",
        "III-1-4：能獨立站立幾秒",
        "III-1-5：能將物品由一手換到另一手",
        "III-1-6：能用拇指配合其他手指鉗握物品",
        "III-1-7：能以雙手拍手",
        "III-2-1：會揮手表示再見",
        "III-2-2：能和成人玩簡單重複遊戲",
        "III-2-3：能簡單表達自我需求及情緒",
        "III-2-4：能區辨照顧者的語氣及情緒",
        "III-2-5：能與照顧者建立情感依附",
        "III-2-6：練習處理陌生人焦慮",
        "III-3-1：能模仿大人的簡單話語",
        "III-3-2：牙牙學語中出現聲量、高低和節奏變化",
        "III-3-3：會以肢體動作進行溝通",
        "III-3-4：會與人輪流對話",
        "III-3-5：能理解簡單語彙的意思",
        "III-4-1：會分辨熟悉家人與陌生人",
        "III-4-2：會尋找完全被藏著的物品",
        "III-4-3：能預期事件的發生",
        "III-4-4：呼叫名字時會有反應",
        "III-4-5：能設法接近想要的事物",
        "III-4-6：能操作簡單玩具",
        "III-5-1：能自己拿住奶瓶進食",
        "III-5-2：能吞嚥糊狀副食品",
        "III-5-3：能自己拿食物吃",
        "III-5-4：能拉下頭上的帽子",
        "III-5-5：會表示要吃東西",
    ],
    "12-24個月": [
        "V-1-1：能穩定移動並調整行走方向",
        "V-1-2：能模仿簡單大肌肉動作",
        "V-1-3：能進行蹲、站、跨等姿勢轉換",
        "V-1-4：能推、踢、搬運大型物件",
        "V-1-5：能進行抓、放、舀、捏等精細操作",
        "V-1-6：能以手指進行較精細的操作",
        "V-2-1：能辨識自己與熟悉人物",
        "V-2-2：能參與簡單團體活動",
        "V-2-3：能辨識或表達基本情緒",
        "V-2-6：能與他人進行簡單社交互動",
        "V-2-8：能模仿照顧、分享等社會性行為",
        "V-3-1：能理解並使用熟悉詞彙",
        "V-3-2：能以簡單語詞或短句表達",
        "V-3-5：能配合童謠、節奏或律動活動",
        "V-3-6：能使用簡單問句或回應問句",
        "V-4-1：能以感官探索物體特徵",
        "V-4-3：能依單一特徵進行簡單分類",
        "V-4-4：能辨識大小、多少等明顯差異",
        "V-4-6：能進行塗鴉、刷畫、拓印等創作探索",
        "V-5-1：能練習自行使用湯匙進食",
        "V-5-2：能練習咀嚼適齡固體食物",
        "V-5-3：能參與簡單穿脫",
        "V-5-6：能參與擦手、擦嘴等清潔",
        "V-5-7：能參與收拾與簡單生活工作",
    ],
    "24-36個月": [
        # 本中心向上延伸能力；若要作為正式評量，請換成中心核定之正式指標。
        "E-1：能跑、跳、跨、繞等方式控制身體移動",
        "E-2：能使用夾子、粗筆、黏土、拼圖等進行精細操作",
        "E-3：能在團體中等待、輪流、分享與合作",
        "E-4：能以簡短句子表達需求、感受與生活經驗",
        "E-5：能依顏色、形狀、大小、用途等特徵分類或配對",
        "E-6：能完成簡單2步驟指令與問題解決",
        "E-7：能主動參與收拾、擦拭、穿脫等生活自理",
        "E-8：能運用多種媒材進行塗畫、撕貼、拓印與立體創作",
        "E-9：能配合音樂節奏進行律動與簡單敲擊",
    ],
}



# ============================================================
# 指標代碼正規化工具
# AI 可回傳「V-1-2」或「V-1-2：完整文字」，系統都能辨識。
# 最後會統一轉回本程式指標池中的完整文字。
# ============================================================

def indicator_code(indicator_text):
    """擷取指標代碼。"""
    if indicator_text is None:
        return ""

    s = str(indicator_text).strip().upper()

    # 支援 I-1-1 / II-3-2 / III-4-6 / IV-1-7 / V-1-2 / VI-4-3 / E-3
    m = re.search(r'\b(?:I{1,3}|IV|V|VI)-\d+(?:-\d+)?\b', s)
    if m:
        return m.group(0)

    return s


def build_indicator_lookup(age_group):
    """建立 {代碼: 完整指標文字} 對照表。"""
    lookup = {}
    for item in INDICATOR_POOLS.get(age_group, []):
        code = indicator_code(item)
        if code:
            lookup[code] = item
    return lookup


def normalize_indicators(age_group, indicators):
    """把 AI 回傳的代碼／完整文字正規化為本地指標池完整文字。"""
    lookup = build_indicator_lookup(age_group)
    normalized = []
    invalid = []

    for raw in indicators or []:
        code = indicator_code(raw)
        if code in lookup:
            normalized.append(lookup[code])
        else:
            invalid.append(str(raw))

    normalized = list(dict.fromkeys(normalized))
    invalid = list(dict.fromkeys(invalid))
    return normalized, invalid


# ============================================================
# OpenAI Structured Outputs Schema
# ============================================================

DOMAIN_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "activity_name": {"type": "string"},
        "activity_1": {"type": "string"},
        "activity_2": {"type": "string"},
        "indicators": {
            "type": "array",
            "items": {"type": "string"},
        },
    },
    "required": ["activity_name", "activity_1", "activity_2", "indicators"],
}

PLAN_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "weeks": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "week_index": {"type": "integer"},
                    "theme": {"type": "string"},
                    "book": {"type": "string"},
                    "body": DOMAIN_SCHEMA,
                    "social": DOMAIN_SCHEMA,
                    "language": DOMAIN_SCHEMA,
                    "cognition": DOMAIN_SCHEMA,
                    "selfcare": DOMAIN_SCHEMA,
                    "art": DOMAIN_SCHEMA,
                    "environment": {"type": "string"},
                },
                "required": [
                    "week_index",
                    "theme",
                    "book",
                    "body",
                    "social",
                    "language",
                    "cognition",
                    "selfcare",
                    "art",
                    "environment",
                ],
            },
        },
        "monthly_environment": {"type": "string"},
    },
    "required": ["weeks", "monthly_environment"],
}



def build_plan_schema_for_age(age_group):
    """
    依目前月齡建立 Structured Outputs Schema。
    indicators 直接使用 enum 限制，因此 AI 不可能回傳其他月齡的 E-1 / V-1-2 等錯誤代碼。
    """
    allowed_codes = [
        indicator_code(item)
        for item in INDICATOR_POOLS.get(age_group, [])
        if indicator_code(item)
    ]

    domain_schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "activity_name": {"type": "string"},
            "activity_1": {"type": "string"},
            "activity_2": {"type": "string"},
            "indicators": {
                "type": "array",
                "items": {
                    "type": "string",
                    "enum": allowed_codes,
                },
                "minItems": 1,
                "maxItems": 2,
            },
        },
        "required": [
            "activity_name",
            "activity_1",
            "activity_2",
            "indicators",
        ],
    }

    return {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "weeks": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "week_index": {"type": "integer"},
                        "theme": {"type": "string"},
                        "book": {"type": "string"},
                        "body": domain_schema,
                        "social": domain_schema,
                        "language": domain_schema,
                        "cognition": domain_schema,
                        "selfcare": domain_schema,
                        "art": domain_schema,
                        "environment": {"type": "string"},
                    },
                    "required": [
                        "week_index",
                        "theme",
                        "book",
                        "body",
                        "social",
                        "language",
                        "cognition",
                        "selfcare",
                        "art",
                        "environment",
                    ],
                },
            },
            "monthly_environment": {"type": "string"},
        },
        "required": ["weeks", "monthly_environment"],
    }


# ============================================================
# AI 生成
# ============================================================

def get_openai_client(api_key):
    """只使用本次網頁輸入的 API Key；不寫入程式、檔案、資料庫或網址。"""
    api_key = (api_key or "").strip()
    if not api_key:
        return None
    return OpenAI(api_key=api_key)


def schedule_signature(branch, year_roc, month, age_group, weekly_schedule):
    raw = json.dumps(
        {
            "branch": branch,
            "year_roc": year_roc,
            "month": month,
            "age_group": age_group,
            "weekly_schedule": weekly_schedule,
        },
        ensure_ascii=False,
        sort_keys=True,
    )
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def generate_ai_month_plan(
    branch,
    year_roc,
    month,
    age_group,
    weekly_schedule,
    api_key,
    model="gpt-5.1",
):
    client = get_openai_client(api_key)
    if client is None:
        raise RuntimeError("請先在網頁輸入 OpenAI API Key。")

    indicators = INDICATOR_POOLS[age_group]

    instructions = """
你是臺灣托嬰中心的教保課程設計助手。請依照使用者提供的：
1. 月齡／年齡組別
2. 每週主題
3. 每週繪本
4. 托嬰中心日常適性活動原則
設計每週六大領域的「適性發展活動月計畫」。

必要規則：
- 每一週都必須真正圍繞該週主題與繪本重新設計，禁止把前一週的素材換書名後重複使用。
- 身體動作、社會情緒、語言溝通、認知探索、生活自理、文化藝術六大領域都要與主題有自然連結。
- 活動必須符合選定年齡，不得出現明顯超齡操作。
- 只能使用使用者提供的指標清單；不得自行發明任何指標代碼。
- 每個領域選1至2個最相關指標即可。
- 活動材料需符合托嬰安全：避免可吞食小物、尖銳物、未密封豆類、窒息風險材料。
- 食物活動需寫成「依中心飲食與過敏規範評估後進行」，不可假設所有幼兒皆能食用。
- 生活自理應融入托育日常，不必勉強與節慶做表面連結。
- 文字使用繁體中文，語氣專業、具體，老師可直接執行。
- 文化藝術重點為感官、塗鴉、拓印、撕貼、律動等歷程，不強調成品一致。
- 24-36個月若使用E系列，請視為中心延伸能力，不寫成政府官方指標。
"""

    user_payload = {
        "園所": branch,
        "民國年": year_roc,
        "月份": month,
        "年齡組別": age_group,
        "可使用指標": indicators,
        "每週資料": weekly_schedule,
    }

    response = client.responses.create(
        model=model,
        instructions=instructions,
        input=json.dumps(user_payload, ensure_ascii=False, indent=2),
        text={
            "format": {
                "type": "json_schema",
                "name": "aosen_monthly_plan",
                "schema": build_plan_schema_for_age(age_group),
                "strict": True,
            }
        },
    )

    if not response.output_text:
        raise RuntimeError("AI沒有回傳可用內容，請重新產生。")

    result = json.loads(response.output_text)

    # 基本防呆：週數必須一致
    if len(result.get("weeks", [])) != len(weekly_schedule):
        raise RuntimeError("AI回傳週數與輪值表不一致，請重新產生。")

    # 再防呆：AI只能使用本地指標池。
    # AI即使只回傳代碼（例如 V-1-2），也會自動轉成完整指標文字。
    for week in result["weeks"]:
        for key in ["body", "social", "language", "cognition", "selfcare", "art"]:
            chosen = week[key].get("indicators", [])
            normalized, invalid = normalize_indicators(age_group, chosen)

            if invalid:
                raise RuntimeError(
                    "AI回傳了不屬於目前月齡的指標："
                    + "、".join(invalid)
                    + "。請重新產生。"
                )

            if not normalized:
                raise RuntimeError(
                    f"AI在「{key}」領域沒有提供可用的適齡指標，請重新產生。"
                )

            # 網頁預覽與 Word 一律使用本地指標池中的完整文字
            week[key]["indicators"] = normalized

    return result


# ============================================================
# Word 1：輪值表
# ============================================================

def generate_roster_docx(
    branch_name,
    year_roc,
    month,
    teachers_str,
    table_data,
    headers,
):
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(3)

    r = title_p.add_run(
        f"{branch_name} {year_roc} 年 {month} 月教案輪值表"
    )
    set_run_font(r, 16, True, (31, 73, 125))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(6)

    sub_text = (
        "幼兒發展領域：身體動作、社會情緒、語言溝通、認知探索、生活自理"
    )
    if teachers_str.strip():
        sub_text += f"\n主帶托育人員：{teachers_str}"

    r = sub_p.add_run(sub_text)
    set_run_font(r, 11)

    table = doc.add_table(rows=len(table_data) + 1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    col_widths = [
        Inches(1.25),
        Inches(2.0),
        Inches(1.8),
        Inches(1.8),
        Inches(1.8),
        Inches(1.8),
    ]

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = col_widths[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_background(cell, "1F497D")
        set_cell_margins(cell, 80, 80, 60, 60)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, 10, True, (255, 255, 255))

    special_keywords = [
        "特約醫師",
        "牙醫師",
        "消防演練",
        "大型活動",
        "國定假日",
        "教案暫停",
        "連假",
        "中秋節",
        "教師節",
        "體能暫停",
    ]

    for row_idx, row_content in enumerate(table_data, start=1):
        row = table.rows[row_idx]
        prevent_row_split(row)

        for col_idx, text in enumerate(row_content):
            cell = table.cell(row_idx, col_idx)
            cell.width = col_widths[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 70, 70, 60, 60)

            if col_idx == 0:
                set_cell_background(cell, "DCE6F1")
            else:
                set_cell_background(
                    cell,
                    "FAFAFA" if row_idx % 2 == 1 else "FFFFFF",
                )

            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if col_idx == 0
                else WD_ALIGN_PARAGRAPH.LEFT
            )

            r = p.add_run(str(text))
            is_special = (
                col_idx > 0
                and any(k in str(text) for k in special_keywords)
            )

            if col_idx == 0:
                set_run_font(r, 11, True, (31, 73, 125))
            elif is_special:
                set_run_font(r, 11, True, (192, 0, 0))
            else:
                set_run_font(r, 11)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# ============================================================
# Word 2：適性月計畫
# ============================================================

def format_domain_for_word(domain_data):
    indicators = "、".join(domain_data.get("indicators", []))
    return (
        f"【{domain_data.get('activity_name', '')}】\n"
        f"1. {domain_data.get('activity_1', '')}\n"
        f"2. {domain_data.get('activity_2', '')}\n"
        f"適齡指標：{indicators}"
    )


def generate_ai_plan_docx(
    branch_name,
    year_roc,
    month,
    age_group,
    teachers_str,
    weekly_schedule,
    plan_data,
):
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.35)
        section.right_margin = Inches(0.35)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)

    r = title_p.add_run(
        f"{branch_name} 適性發展活動 {year_roc} 年 {month} 月計畫"
    )
    set_run_font(r, 15, True, (31, 73, 125))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(3)

    r = sub_p.add_run(
        f"月份：{year_roc} 年 {month} 月　｜　"
        f"年齡/組別：{age_group}　｜　"
        f"主責托育人員：{teachers_str}"
    )
    set_run_font(r, 9.5, True)

    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_after = Pt(3)
    r = note_p.add_run(
        "填表原則：依嬰幼兒個別差異、興趣及當日身心狀況彈性調整；"
        "活動融入日常作息中重複進行，不作為單一發展評量。"
    )
    set_run_font(r, 8)

    table = doc.add_table(
        rows=len(weekly_schedule) + 1,
        cols=7,
    )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    headers = [
        "週次／主題",
        "身體動作",
        "社會情緒",
        "語言溝通",
        "認知探索",
        "生活自理",
        "文化藝術",
    ]

    widths = [
        Inches(1.35),
        Inches(1.55),
        Inches(1.55),
        Inches(1.55),
        Inches(1.55),
        Inches(1.55),
        Inches(1.55),
    ]

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = widths[i]
        set_cell_background(cell, "1F497D")
        set_cell_margins(cell, 40, 40, 40, 40)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, 8.5, True, (255, 255, 255))

    ai_weeks = plan_data["weeks"]

    for idx, schedule_item in enumerate(weekly_schedule):
        ai_week = ai_weeks[idx]
        row = table.rows[idx + 1]
        prevent_row_split(row)

        week_title = (
            f"{schedule_item['week_range']}\n"
            f"【主題：{schedule_item['theme']}】\n"
            f"繪本：《{schedule_item['book']}》\n"
            f"主責：{schedule_item['roster_summary']}"
        )

        row_content = [
            week_title,
            format_domain_for_word(ai_week["body"]),
            format_domain_for_word(ai_week["social"]),
            format_domain_for_word(ai_week["language"]),
            format_domain_for_word(ai_week["cognition"]),
            format_domain_for_word(ai_week["selfcare"]),
            format_domain_for_word(ai_week["art"]),
        ]

        for col_idx, text in enumerate(row_content):
            cell = table.cell(idx + 1, col_idx)
            cell.width = widths[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell, 35, 35, 35, 35)

            if col_idx == 0:
                set_cell_background(cell, "DCE6F1")
            else:
                set_cell_background(
                    cell,
                    "FAFAFA" if idx % 2 == 0 else "FFFFFF",
                )

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 0.9

            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            r = p.add_run(text)
            set_run_font(
                r,
                7.4 if col_idx > 0 else 7.8,
                col_idx == 0,
                (31, 73, 125) if col_idx == 0 else None,
            )

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    env_table = doc.add_table(rows=2, cols=1)
    env_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(env_table)

    cell = env_table.cell(0, 0)
    set_cell_background(cell, "1F497D")
    p = cell.paragraphs[0]
    r = p.add_run("教材、情境、環境調整與教玩具規劃")
    set_run_font(r, 8.5, True, (255, 255, 255))

    cell = env_table.cell(1, 0)
    set_cell_background(cell, "FAFAFA")
    set_cell_margins(cell, 40, 40, 60, 60)

    p = cell.paragraphs[0]
    env_lines = []
    for idx, schedule_item in enumerate(weekly_schedule):
        env_lines.append(
            f"{schedule_item['week_range']}／{schedule_item['theme']}："
            f"{plan_data['weeks'][idx]['environment']}"
        )

    env_lines.append(
        f"整月環境規劃：{plan_data.get('monthly_environment', '')}"
    )

    r = p.add_run("\n".join(env_lines))
    set_run_font(r, 7.6)

    if age_group == "24-36個月":
        p = doc.add_paragraph()
        r = p.add_run(
            "註：24–36個月之E系列為本中心延伸能力描述，"
            "不標示為政府官方發展評量指標。"
        )
        set_run_font(r, 7.5)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# ============================================================
# UI
# ============================================================

st.title("🏫 托嬰中心 教案與適性月計畫")

st.markdown(
    "完成輪值排班後，可下載輪值表；"
    "再選擇月齡，依每週主題與繪本重新設計適性月計畫。"
)

# ============================================================
# 壹、基本設定
# ============================================================

st.subheader("壹、基本設定")

c1, c2, c3 = st.columns([1.5, 1, 1])

with c1:
    branch = st.selectbox(
        "園所：",
        ["西湖", "文德"],
    )

with c2:
    year_roc = st.number_input(
        "民國年份：",
        min_value=114,
        max_value=125,
        value=115,
    )

with c3:
    month = st.selectbox(
        "月份：",
        list(range(1, 13)),
        index=8,
    )

teachers_input = st.text_input(
    "主帶老師名單（請以逗號分隔）：",
    value="",
    placeholder="例如：A, B, C",
)

st.markdown("#### 📖 使用說明")
st.markdown(
    """
1. 選取園所、年份及月份  
2. 輸入該月設計教案的人員  
3. 填寫輪值排班表  
4. 完成後檢查，即可一鍵下載輪值表及月計畫
"""
)

builtin_options = [
    "（請選擇）",
    "主任",
    "秋馨", 
    "特約醫師",
    "牙醫師",
    "消防演練",
    "大型活動",
    "國定假日",
    "教案暫停",
]

raw_teachers = [
    t.strip()
    for t in teachers_input.split(",")
    if t.strip()
]

dropdown_options = (
    builtin_options[:2]
    + raw_teachers
    + builtin_options[2:]
)

if branch == "西湖":
    headers = [
        "主題",
        "星期一：繪本",
        "星期二：小肌肉／認知",
        "星期三：小肌肉／觸覺",
        "星期四：體能課",
        "星期五：藝術創作",
    ]
    wed_title_prefix = "週三：小肌肉／觸覺"
    thu_title_prefix = "週四：體能"
else:
    headers = [
        "主題",
        "星期一：繪本",
        "星期二：小肌肉／認知",
        "星期三：體能課",
        "星期四：小肌肉／觸覺",
        "星期五：藝術創作",
    ]
    wed_title_prefix = "週三：體能"
    thu_title_prefix = "週四：小肌肉／觸覺"


# ============================================================
# 貳、輪值排班
# ============================================================

st.divider()
st.subheader(
    f"貳、{branch} {year_roc} 年 {month} 月輪值排班"
)

year_ad = year_roc + 1911
cal = calendar.monthcalendar(year_ad, month)

work_weeks = []
for week in cal:
    mon_to_fri = week[0:5]
    if any(day > 0 for day in mon_to_fri):
        work_weeks.append(mon_to_fri)

table_data = []
weekly_schedule_for_plan = []

for idx, week in enumerate(work_weeks, start=1):

    with st.expander(
        f"📌 第 {idx} 週排班設定",
        expanded=True,
    ):

        cols = st.columns(
            [1.2, 2.2, 1.3, 1.3, 1.3, 1.3]
        )
        col_theme, col_m, col_t, col_w, col_th, col_f = cols

        d_mon = f"{month}/{week[0]}" if week[0] > 0 else ""
        d_tue = f"{month}/{week[1]}" if week[1] > 0 else ""
        d_wed = f"{month}/{week[2]}" if week[2] > 0 else ""
        d_thu = f"{month}/{week[3]}" if week[3] > 0 else ""
        d_fri = f"{month}/{week[4]}" if week[4] > 0 else ""

        valid_days = [d for d in week if d > 0]

        if valid_days:
            w_start = f"{month:02d}/{valid_days[0]:02d}"
            w_end = f"{month:02d}/{valid_days[-1]:02d}"
            week_range_str = (
                f"第{idx}週 ({w_start}-{w_end})"
            )
        else:
            week_range_str = f"第{idx}週"

        if idx == 1:
            theme_default = "顏色"
            book_name_default = "小藍與小黃"
            m_lead_default_idx = 1
        else:
            theme_default = ""
            book_name_default = ""
            m_lead_default_idx = 0

        with col_theme:
            theme = st.text_input(
                "週主題",
                key=f"theme_{idx}",
                value=theme_default,
                placeholder="請輸入主題",
            )

        with col_m:
            st.markdown(
                f"**週一：繪本 ({d_mon})**"
                if d_mon
                else "**週一：繪本**"
            )

            m_lead = st.selectbox(
                "導讀人",
                dropdown_options,
                key=f"m_lead_{idx}",
                index=m_lead_default_idx,
            )

            book_name = st.text_input(
                "繪本名稱",
                key=f"book_{idx}",
                value=book_name_default,
                placeholder="請輸入繪本名稱",
            )

            lead_txt = (
                ""
                if m_lead == "（請選擇）"
                else f" {m_lead}"
            )

            book_txt = (
                f"\n{book_name}"
                if book_name.strip()
                else ""
            )

            mon_str = (
                f"{d_mon}{lead_txt}{book_txt}".strip()
                if d_mon
                else ""
            )

        with col_t:
            st.markdown(
                f"**週二：小肌肉／認知 ({d_tue})**"
                if d_tue
                else "**週二：小肌肉／認知**"
            )

            tue_teacher = st.selectbox(
                "負責人／狀態",
                dropdown_options,
                key=f"tue_{idx}",
                index=0,
            )

            tue_txt = (
                ""
                if tue_teacher == "（請選擇）"
                else f" {tue_teacher}"
            )

            tue_str = (
                f"{d_tue}{tue_txt}".strip()
                if d_tue
                else ""
            )

        with col_w:
            st.markdown(
                f"**{wed_title_prefix} ({d_wed})**"
                if d_wed
                else f"**{wed_title_prefix}**"
            )

            wed_teacher = st.selectbox(
                "負責人／狀態",
                dropdown_options,
                key=f"wed_{idx}",
                index=0,
            )

            wed_txt = (
                ""
                if wed_teacher == "（請選擇）"
                else f" {wed_teacher}"
            )

            wed_str = (
                f"{d_wed}{wed_txt}".strip()
                if d_wed
                else ""
            )

        with col_th:
            st.markdown(
                f"**{thu_title_prefix} ({d_thu})**"
                if d_thu
                else f"**{thu_title_prefix}**"
            )

            thu_teacher = st.selectbox(
                "負責人／狀態",
                dropdown_options,
                key=f"thu_{idx}",
                index=0,
            )

            thu_txt = (
                ""
                if thu_teacher == "（請選擇）"
                else f" {thu_teacher}"
            )

            thu_str = (
                f"{d_thu}{thu_txt}".strip()
                if d_thu
                else ""
            )

        with col_f:
            st.markdown(
                f"**週五：藝術創作 ({d_fri})**"
                if d_fri
                else "**週五：藝術創作**"
            )

            fri_teacher = st.selectbox(
                "負責人／狀態",
                dropdown_options,
                key=f"fri_{idx}",
                index=0,
            )

            fri_txt = (
                ""
                if fri_teacher == "（請選擇）"
                else f" {fri_teacher}"
            )

            fri_str = (
                f"{d_fri}{fri_txt}".strip()
                if d_fri
                else ""
            )

        table_data.append(
            [
                theme,
                mon_str,
                tue_str,
                wed_str,
                thu_str,
                fri_str,
            ]
        )

        assigned_teachers = [
            x
            for x in [
                m_lead,
                tue_teacher,
                wed_teacher,
                thu_teacher,
                fri_teacher,
            ]
            if x not in [
                "（請選擇）",
                "教案暫停",
                "國定假日",
            ]
        ]

        assigned_teachers = list(
            dict.fromkeys(assigned_teachers)
        )

        roster_summary = (
            "、".join(assigned_teachers)
            if assigned_teachers
            else "全體托育人員"
        )

        weekly_schedule_for_plan.append(
            {
                "week_index": idx,
                "week_range": week_range_str,
                "theme": (
                    theme.strip()
                    if theme.strip()
                    else "主題探索"
                ),
                "book": (
                    book_name.strip()
                    if book_name.strip()
                    else "主題繪本導讀"
                ),
                "roster_summary": roster_summary,
            }
        )


# ============================================================
# 參、成果匯出
# ============================================================

st.divider()
st.subheader("參、成果匯出")

st.markdown(
    f"### 📋 1. 下載【{branch}月教案輪值表】"
)

doc_roster_bytes = generate_roster_docx(
    branch,
    year_roc,
    month,
    teachers_input,
    table_data,
    headers,
)

st.download_button(
    label=(
        f"📥 下載【{branch} "
        f"{year_roc}年{month}月教案輪值表】"
    ),
    data=doc_roster_bytes,
    file_name=(
        f"{branch}_"
        f"{year_roc}年{month}月教案輪值表.docx"
    ),
    mime=(
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    ),
    use_container_width=True,
)

st.markdown("---")
st.markdown(
    f"### 🌱 2. 產生【{branch}適性發展活動月計畫】"
)

if branch == "西湖":
    available_ages = [
        "0-12個月",
        "12-24個月",
        "24-36個月",
    ]
else:
    available_ages = [
        "7-12個月",
        "12-24個月",
        "24-36個月",
    ]

age_group = st.selectbox(
    "月計畫年齡階段：",
    available_ages,
    key="month_plan_age_group",
)

st.markdown("#### 🔑 連線設定")

api_key_input = st.text_input(
    "OpenAI API Key：",
    type="password",
    placeholder="請貼上你的 OpenAI API Key（例如 sk-...）",
    help=(
        "此欄位只用於本次瀏覽器工作階段呼叫 OpenAI API。"
        "程式不會把 Key 寫入 Word、GitHub、資料庫或網址參數。"
    ),
)

model_name = st.selectbox(
    "AI 模型：",
    ["gpt-5-mini", "gpt-4.1-mini"],
    index=0,
    help="一般建議使用gpt-5-mini。",
)

st.caption(
    "🔒 API Key 僅在你按下「產生／更新適性月計畫」時，"
    "由 Streamlit 伺服器暫時用於呼叫 OpenAI；"
    "不會寫入程式碼、下載文件、GitHub、資料庫或網址。"
)

current_sig = schedule_signature(
    branch,
    year_roc,
    month,
    age_group,
    weekly_schedule_for_plan,
)

if "generated_plan" not in st.session_state:
    st.session_state.generated_plan = None

if "generated_plan_signature" not in st.session_state:
    st.session_state.generated_plan_signature = None

generate_col, clear_col = st.columns([3, 1])

with generate_col:
    if st.button(
        "✨ 產生／更新適性月計畫",
        type="primary",
        use_container_width=True,
    ):
        try:
            with st.spinner(
                "正在依每週主題、繪本與適齡指標設計月計畫..."
            ):
                if not api_key_input.strip():
                    raise RuntimeError("請先輸入 OpenAI API Key。")

                result = generate_ai_month_plan(
                    branch,
                    year_roc,
                    month,
                    age_group,
                    weekly_schedule_for_plan,
                    api_key_input,
                    model_name,
                )
                st.session_state.generated_plan = result
                st.session_state.generated_plan_signature = current_sig
                st.success(
                    "適性月計畫已完成，請先預覽內容再下載。"
                )
        except Exception as e:
            st.error(f"產生失敗：{e}")

with clear_col:
    if st.button(
        "清除內容",
        use_container_width=True,
    ):
        st.session_state.generated_plan = None
        st.session_state.generated_plan_signature = None
        st.rerun()

if (
    st.session_state.generated_plan is not None
    and st.session_state.generated_plan_signature != current_sig
):
    st.warning(
        "你已修改園所、月份、月齡、主題或繪本。"
        "目前預覽是舊版本，請重新按「產生／更新適性月計畫」。"
    )

plan_data = st.session_state.generated_plan

if plan_data is not None:
    st.markdown("#### 👀 月計畫預覽")

    domain_labels = {
        "body": "身體動作",
        "social": "社會情緒",
        "language": "語言溝通",
        "cognition": "認知探索",
        "selfcare": "生活自理",
        "art": "文化藝術",
    }

    for idx, week in enumerate(plan_data["weeks"]):
        schedule_item = weekly_schedule_for_plan[idx]

        with st.expander(
            f"{schedule_item['week_range']}｜"
            f"{schedule_item['theme']}｜"
            f"《{schedule_item['book']}》",
            expanded=False,
        ):
            for key, label in domain_labels.items():
                d = week[key]
                st.markdown(
                    f"**{label}｜{d['activity_name']}**"
                )
                st.write(f"1. {d['activity_1']}")
                st.write(f"2. {d['activity_2']}")
                st.caption(
                    "適齡指標："
                    + "、".join(d["indicators"])
                )

            st.caption(
                "環境與材料："
                + week["environment"]
            )

    if st.session_state.generated_plan_signature == current_sig:
        ai_doc_bytes = generate_ai_plan_docx(
            branch,
            year_roc,
            month,
            age_group,
            teachers_input,
            weekly_schedule_for_plan,
            plan_data,
        )

        st.download_button(
            label=(
                f"📥 下載【{branch}（{age_group}）"
                f"{year_roc}年{month}月適性月計畫】"
            ),
            data=ai_doc_bytes,
            file_name=(
                f"{branch}_"
                f"{year_roc}年{month}月_"
                f"{age_group}_適性發展活動月計畫.docx"
            ),
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            use_container_width=True,
        )
